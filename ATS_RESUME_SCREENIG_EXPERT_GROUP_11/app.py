from dotenv import load_dotenv
import base64
import streamlit as st
import os
import io
from PIL import Image
import pdf2image
import google.generativeai as genai
import tempfile
import PyPDF2
import time
from docx import Document
import logging
from typing import List, Dict, Optional, Tuple, Any

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
PAGE_TITLE = "ATS Resume Expert"
PAGE_ICON = "📄"
LOGO_PATH = "assets/logo.png"
CSS_PATH = "assets/styles.css"
MODEL_NAME = "gemini-1.5-flash"
FEEDBACK_FILE = "feedback_log.txt"

class ATSResumeExpert:
    """Main application class for ATS Resume Expert"""
    
    def __init__(self):
        """Initialize the application"""
        self.api_key = os.getenv("GEMINI_API_KEY")
        self.setup_page_config()
        self.load_css()
        
    def setup_page_config(self):
        """Configure Streamlit page settings"""
        st.set_page_config(
            page_title=PAGE_TITLE,
            page_icon=PAGE_ICON,
            layout="wide"
        )
    
    def load_css(self):
        """Load custom CSS styles"""
        try:
            with open(CSS_PATH) as f:
                st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
        except FileNotFoundError:
            logger.warning(f"CSS file not found at {CSS_PATH}")
    
    def render_sidebar(self):
        """Render sidebar with app information and API key input"""
        with st.sidebar:
            st.image(LOGO_PATH, width=250)
            st.title("📄 ATS Resume Expert")
            st.markdown("""
            **Features:**  
            ✅ Analyze resume vs job description  
            ✅ Get feedback & improvement tips  
            ✅ See keyword match percentage  
            ✅ Download detailed analysis  
            """)
            st.info("🚀 Powered by Google Gemini AI")
            
            if not self.api_key:
                user_api_key = st.text_input("🔑 Enter your Gemini API Key:", type="password")
                if user_api_key:
                    self.api_key = user_api_key
                    return True
            return bool(self.api_key)
    
    def configure_genai(self) -> bool:
        """Configure the Gemini API with the API key"""
        if not self.api_key:
            st.warning("⚠️ Please enter your Gemini API Key to continue.")
            return False
            
        try:
            genai.configure(api_key=self.api_key)
            return True
        except Exception as e:
            logger.error(f"API Configuration Failed: {e}")
            st.error(f"❌ API Configuration Failed: {str(e)}")
            return False
    
    @staticmethod
    def extract_text_from_pdf(uploaded_file) -> str:
        """Extract text content from uploaded PDF file"""
        if not uploaded_file:
            return ""
            
        text_content = ""
        tmp_path = None
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            with open(tmp_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only add non-empty pages
                        text_content += page_text + "\n\n"
            
            return text_content
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            st.error(f"❌ Error extracting text from PDF: {str(e)}")
            return ""
        finally:
            # Clean up temp file
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file: {e}")
    
    @staticmethod
    def convert_pdf_to_images(uploaded_file, page_limit=2) -> Optional[List[Dict]]:
        """Convert PDF to encoded images for AI processing"""
        if not uploaded_file:
            return None
            
        try:
            images = pdf2image.convert_from_bytes(uploaded_file.read(), last_page=page_limit)
            if not images:
                return None
                
            encoded_images = []
            for img in images:
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=85)
                encoded_images.append({
                    "mime_type": "image/jpeg",
                    "data": base64.b64encode(img_byte_arr.getvalue()).decode()
                })
            return encoded_images
        except Exception as e:
            logger.error(f"PDF to image conversion error: {e}")
            st.error(f"❌ Error converting PDF to images: {str(e)}")
            return None
    
    def get_ai_response(self, job_description: str, pdf_images: List[Dict], prompt: str, resume_text: str) -> str:
        """Get AI analysis response from Gemini"""
        if not self.configure_genai():
            return "⚠️ API Key is missing or invalid. Please provide a valid Gemini API key."
            
        try:
            model = genai.GenerativeModel(MODEL_NAME)
            
            # Prepare content for AI
            content = [prompt]
            if job_description:
                content.append(f"**JOB DESCRIPTION:**\n{job_description}")
            if resume_text:
                content.append(f"**RESUME TEXT:**\n{resume_text}")
            if pdf_images:
                content.extend(pdf_images[:2])  # Limited to first 2 pages
                
            # Add a timeout to prevent hanging
            start_time = time.time()
            response = model.generate_content(content)
            generation_time = time.time() - start_time
            logger.info(f"AI response generated in {generation_time:.2f} seconds")
            
            return response.text if hasattr(response, 'text') else "⚠️ No response received from AI."
        except Exception as e:
            logger.error(f"AI processing error: {e}")
            return f"❌ AI Processing Error: {str(e)}"
    
    @staticmethod
    def generate_report_docx(title: str, content: str) -> io.BytesIO:
        """Generate a downloadable DOCX report"""
        doc = Document()
        doc.add_heading(title, level=1)
        
        # Split content by paragraphs and add to document
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    @staticmethod
    def save_feedback(email: str, feedback: str) -> bool:
        """Save user feedback to file"""
        if not feedback.strip():
            return False
            
        try:
            timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
            with open(FEEDBACK_FILE, "a") as file:
                file.write(f"Timestamp: {timestamp}\n")
                file.write(f"Email: {email or 'Anonymous'}\n")
                file.write(f"Feedback: {feedback}\n")
                file.write("-" * 50 + "\n\n")
            return True
        except Exception as e:
            logger.error(f"Error saving feedback: {e}")
            return False
    
    def render_main_ui(self):
        """Render the main UI components"""
        st.title("🚀 ATS Resume Analyzer")
        
        # Create two columns layout
        col1, col2 = st.columns(2)
        
        # Job Description input
        with col1:
            st.subheader("📌 Job Description")
            job_description = st.text_area(
                "Paste the job description here:", 
                height=300,
                help="Copy and paste the full job description for best analysis results"
            )
        
        # Resume upload
        with col2:
            st.subheader("📄 Your Resume")
            uploaded_file = st.file_uploader(
                "Upload your resume (PDF)...", 
                type=["pdf"],
                help="Please upload a PDF document containing your resume"
            )
            
            # Display resume preview
            if uploaded_file:
                with st.expander("👀 Resume Preview", expanded=True):
                    try:
                        pdf_images = pdf2image.convert_from_bytes(uploaded_file.getvalue(), last_page=1)
                        if pdf_images:
                            st.image(pdf_images[0], caption="📄 Resume Preview (First Page)", width=600)
                            st.success("✅ PDF Uploaded Successfully")
                    except Exception as e:
                        st.error(f"❌ Error displaying PDF preview: {str(e)}")
        
        return job_description, uploaded_file
    
    def render_analysis_options(self) -> Dict[str, Tuple[bool, str]]:
        """Render analysis option buttons and return selected options with prompts"""
        st.subheader("🔍 Analysis Options")
        
        col_buttons = st.columns(3)
        
        # Analysis option buttons
        with col_buttons[0]:
            resume_eval = st.button("📋 Resume Evaluation")
        
        with col_buttons[1]:
            keyword_analysis = st.button("🔎 Keyword Analysis")
        
        with col_buttons[2]:
            match_percentage = st.button("📊 Match Percentage")
        
        # Map buttons to prompts
        return {
            "Resume Evaluation": (
                resume_eval, 
                "You are a Technical HR expert. Evaluate this resume against the job description. Provide detailed feedback on strengths, weaknesses, and specific improvements."
            ),
            "Keyword Analysis": (
                keyword_analysis, 
                "You are an ATS keyword specialist. List all matching keywords found in both the resume and job description. Then list important missing keywords that should be included. Format as two separate lists with headings."
            ),
            "Match Percentage": (
                match_percentage, 
                "You are an ATS scanner. Calculate an estimated match percentage between the resume and job description. Explain your reasoning, then provide 3-5 specific recommendations to improve the match score."
            )
        }
    
    def render_feedback_section(self):
        """Render the feedback form section"""
        st.subheader("📝 Your Feedback")
        
        col_f1, col_f2 = st.columns([1, 2])
        
        with col_f1:
            user_email = st.text_input(
                "📧 Email (Optional)", 
                placeholder="your-email@example.com"
            )
        
        feedback_text = st.text_area(
            "💬 Share your experience...", 
            placeholder="Your feedback helps us improve! Let us know what you think of the tool."
        )
        
        if st.button("📤 Submit Feedback"):
            if feedback_text.strip():
                if self.save_feedback(user_email, feedback_text):
                    st.success("✅ Thank you for your feedback!")
                else:
                    st.error("❌ Failed to save feedback. Please try again.")
            else:
                st.warning("⚠️ Please enter some feedback before submitting.")
    
    def run(self):
        """Run the application"""
        # Render sidebar and check API key
        has_api_key = self.render_sidebar()
        
        # Render main UI components
        job_description, uploaded_file = self.render_main_ui()
        
        # Render analysis options
        analysis_options = self.render_analysis_options()
        
        # Process analysis if file uploaded and API key available
        if uploaded_file and has_api_key:
            for option_name, (button_pressed, prompt) in analysis_options.items():
                if button_pressed:
                    with st.spinner(f"⏳ Processing {option_name.lower()}..."):
                        # Extract text from the PDF
                        resume_text = self.extract_text_from_pdf(uploaded_file)
                        
                        # Reset file position after reading
                        uploaded_file.seek(0)
                        
                        # Convert PDF to images for visual analysis
                        pdf_images = self.convert_pdf_to_images(uploaded_file)
                        
                        # Get AI analysis
                        response = self.get_ai_response(job_description, pdf_images, prompt, resume_text)
                        
                        # Display results
                        st.subheader(f"📊 {option_name} Results")
                        st.markdown(response)
                        
                        # Generate downloadable report
                        docx_file = self.generate_report_docx(option_name, response)
                        
                        # Add download button
                        st.download_button(
                            label="📥 Download Report (.docx)",
                            data=docx_file,
                            file_name=f"{option_name.replace(' ', '_').lower()}_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
        elif uploaded_file and not has_api_key:
            st.warning("⚠️ Please enter your Gemini API Key in the sidebar to continue.")
        elif any(button for button, _ in analysis_options.values()) and not uploaded_file:
            st.warning("⚠️ Please upload your resume before analyzing.")
        
        # Render feedback section
        self.render_feedback_section()
        
        # Add footer
        st.markdown("---")
        st.markdown("Made with ❤️ for job seekers | © 2025 ATS Resume Expert")


if __name__ == "__main__":
    load_dotenv()  # Load environment variables
    app = ATSResumeExpert()
    app.run()